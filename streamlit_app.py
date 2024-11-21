import streamlit as st
import PyPDF2
from pdf2image import convert_from_path
from PIL import Image
import pdf2docx
import os
import zipfile
from fpdf import FPDF
import io
import shutil
import requests
import pandas as pd
import json
import markdown
from pydub import AudioSegment
import docx
from io import BytesIO
from pydub.playback import play
from gtts import gTTS
import qrcode

# Helper function to save uploaded files
def save_uploadedfile(uploadedfile, filename="temp.pdf"):
    with open(filename, "wb") as f:
        f.write(uploadedfile.getbuffer())

# PDF Merger
def merge_pdfs(pdf_files):
    merger = PyPDF2.PdfMerger()
    for file in pdf_files:
        try:
            save_uploadedfile(file, "temp.pdf")
            merger.append("temp.pdf")
        except Exception as e:
            st.error(f"Error reading PDF file: {e}")
    with open("merged.pdf", "wb") as merged_pdf:
        merger.write(merged_pdf)
    os.remove("temp.pdf")
    return "merged.pdf"

# PDF Splitter
def split_pdf(pdf_file, page_range):
    try:
        save_uploadedfile(pdf_file, "temp.pdf")
        pdf = PyPDF2.PdfReader("temp.pdf")
        writer = PyPDF2.PdfWriter()
        for page_num in range(page_range[0] - 1, page_range[1]):
            writer.add_page(pdf.pages[page_num])
        with open("split.pdf", "wb") as output_pdf:
            writer.write(output_pdf)
        os.remove("temp.pdf")
        return "split.pdf"
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")

# PDF to Word Converter
def pdf_to_word(pdf_file):
    save_uploadedfile(pdf_file, "temp.pdf")
    converter = pdf2docx.Converter("temp.pdf")
    converter.convert("document.docx")
    converter.close()
    os.remove("temp.pdf")
    return "document.docx"

# Extract Images from PDF
def extract_images_from_pdf(pdf_file):
    save_uploadedfile(pdf_file, "temp.pdf")
    images = convert_from_path("temp.pdf")
    image_files = []
    for i, image in enumerate(images):
        img_filename = f"page_{i+1}.png"
        image.save(img_filename, "PNG")
        image_files.append(img_filename)
    os.remove("temp.pdf")
    return image_files

# Image Manipulation (resize)
def resize_image(image_file, width=800, height=600):
    img = Image.open(image_file)
    img_resized = img.resize((width, height))
    resized_filename = f"resized_{os.path.basename(image_file)}"
    img_resized.save(resized_filename)
    return resized_filename

# Word to PDF (Batch conversion)
def batch_convert_word_to_pdf(word_files):
    pdf_files = []
    for word_file in word_files:
        doc = docx.Document(word_file)
        pdf_filename = f"{os.path.splitext(word_file.name)[0]}.pdf"
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for para in doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)
        pdf.output(pdf_filename)
        pdf_files.append(pdf_filename)
    return pdf_files

# Convert Text from PDF to Speech
def text_to_speech(text):
    tts = gTTS(text, lang='en')
    tts.save("text_to_speech.mp3")
    return "text_to_speech.mp3"

# Audio to MP3 and Playback
def audio_to_mp3(audio_file):
    audio = AudioSegment.from_file(audio_file)
    mp3_filename = "audio_to_mp3.mp3"
    audio.export(mp3_filename, format="mp3")
    return mp3_filename

# Convert Markdown to PDF
def markdown_to_pdf(markdown_file):
    html_content = markdown.markdown(markdown_file.read().decode("utf-8"))
    pdf_filename = "markdown_to_pdf.pdf"
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, html_content)
    pdf.output(pdf_filename)
    return pdf_filename

# Convert ZIP file
def extract_zip(zip_file):
    with zipfile.ZipFile(zip_file, 'r') as zip_ref:
        zip_ref.extractall("extracted_files")
    return "extracted_files"

# Convert Text to PDF
def text_to_pdf(text):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    pdf_filename = "text_to_pdf.pdf"
    pdf.output(pdf_filename)
    return pdf_filename

# Generate QR Code from Text
def generate_qr_code(text):
    qr = qrcode.make(text)
    qr_filename = "qr_code.png"
    qr.save(qr_filename)
    return qr_filename

# Convert Image to PDF
def image_to_pdf(image_file):
    img = Image.open(image_file)
    pdf_filename = f"{os.path.splitext(image_file.name)[0]}.pdf"
    img.convert('RGB').save(pdf_filename)
    return pdf_filename

# PDF Page Rotator
def rotate_pdf(pdf_file, page_num, angle=90):
    save_uploadedfile(pdf_file, "temp.pdf")
    pdf = PyPDF2.PdfReader("temp.pdf")
    writer = PyPDF2.PdfWriter()
    for i, page in enumerate(pdf.pages):
        if i == page_num - 1:
            page.rotate_clockwise(angle)
        writer.add_page(page)
    with open("rotated.pdf", "wb") as rotated_pdf:
        writer.write(rotated_pdf)
    os.remove("temp.pdf")
    return "rotated.pdf"

# File Compression (Zip files)
def compress_files_to_zip(files):
    zip_filename = "compressed_files.zip"
    with zipfile.ZipFile(zip_filename, 'w') as zipf:
        for file in files:
            zipf.write(file, os.path.basename(file))
    return zip_filename

# Text File to Word Converter
def text_to_word(text_file):
    with open(text_file, "r") as f:
        content = f.read()
    doc = docx.Document()
    doc.add_paragraph(content)
    word_filename = "document_from_text.docx"
    doc.save(word_filename)
    return word_filename

# Streamlit App
st.title("Advanced File Conversion Tools")

# PDF Merger
st.header("PDF Merger")
pdf_files = st.file_uploader("Select PDF files to merge", type=["pdf"], accept_multiple_files=True, key="merge_uploader")
if st.button("Merge PDFs"):
    if pdf_files:
        merged_pdf = merge_pdfs(pdf_files)
        st.download_button("Download Merged PDF", merged_pdf, file_name="merged.pdf")

# PDF to Word Converter
st.header("PDF to Word Converter")
pdf_file = st.file_uploader("Select PDF file to convert", type=["pdf"], key="word_uploader")
if st.button("Convert to Word"):
    if pdf_file:
        docx_file = pdf_to_word(pdf_file)
        st.download_button("Download Word Document", docx_file, file_name="document.docx")

# PDF to Text-to-Speech
st.header("Text-to-Speech from PDF")
pdf_text_file = st.file_uploader("Select PDF file for text-to-speech", type=["pdf"], key="text_to_speech_uploader")
if st.button("Convert to Speech"):
    if pdf_text_file:
        text = extract_text(pdf_text_file)
        audio_file = text_to_speech(text)
        st.download_button("Download Speech MP3", audio_file, file_name="text_to_speech.mp3")
        audio = AudioSegment.from_mp3(audio_file)
        play(audio)

# Extract Images from PDF
st.header("Extract Images from PDF")
pdf_image_file = st.file_uploader("Select PDF file to extract images", type=["pdf"], key="image_extractor_uploader")
if st.button("Extract Images"):
    if pdf_image_file:
        images = extract_images_from_pdf(pdf_image_file)
        for img in images:
            st.image(img)

# Image Manipulation (Resize)
st.header("Resize Image Before Conversion")
image_file = st.file_uploader("Select Image to Resize", type=["jpg", "jpeg", "png"], key="resize_uploader")
width = st.number_input("Width (px)", min_value=100, max_value=2000, value=800)
height = st.number_input("Height (px)", min_value=100, max_value=2000, value=600)
if st.button("Resize Image"):
    if image_file:
        resized_image = resize_image(image_file, width, height)
        st.image(resized_image)

# Word to PDF (Batch conversion)
st.header("Batch Convert Word to PDF")
word_files = st.file_uploader("Select Word files to convert", type=["docx"], accept_multiple_files=True, key="batch_word_uploader")
if st.button("Batch Convert to PDF"):
    if word_files:
        pdf_files = batch_convert_word_to_pdf(word_files)
        for pdf_file in pdf_files:
            st.download_button(f"Download {os.path.basename(pdf_file)}", pdf_file, file_name=pdf_file)

# ZIP File Extraction
st.header("Extract ZIP File")
zip_file = st.file_uploader("Select ZIP file to extract", type=["zip"], key="zip_uploader")
if st.button("Extract ZIP"):
    if zip_file:
        extracted_files = extract_zip(zip_file)
        st.write(f"Files extracted to: {extracted_files}")

# Audio to MP3
st.header("Audio to MP3 Converter")
audio_file = st.file_uploader("Select Audio to Convert", type=["wav", "flac", "ogg"], key="audio_uploader")
if st.button("Convert to MP3"):
    if audio_file:
        mp3_file = audio_to_mp3(audio_file)
        st.download_button("Download MP3", mp3_file, file_name="audio.mp3")

# Generate QR Code
st.header("Generate QR Code")
qr_text = st.text_input("Enter text for QR code")
if st.button("Generate QR Code"):
    if qr_text:
        qr_code_file = generate_qr_code(qr_text)
        st.image(qr_code_file)

# Image to PDF
st.header("Image to PDF")
image_file_pdf = st.file_uploader("Select Image(s) to Convert to PDF", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key="image_pdf_uploader")
if st.button("Convert Image(s) to PDF"):
    if image_file_pdf:
        pdf_file = image_to_pdf(image_file_pdf[0])  # Converts the first image file
        st.download_button("Download PDF", pdf_file)

# Compress Files to ZIP
st.header("Compress Files to ZIP")
files_to_compress = st.file_uploader("Select files to compress", type=["txt", "jpg", "pdf", "docx"], accept_multiple_files=True, key="compress_uploader")
if st.button("Compress Files"):
    if files_to_compress:
        compressed_file = compress_files_to_zip(files_to_compress)
        st.download_button("Download Compressed ZIP", compressed_file)

# Convert Text File to Word
st.header("Text File to Word Converter")
text_file = st.file_uploader("Select Text File to Convert to Word", type=["txt"], key="text_to_word_uploader")
if st.button("Convert to Word"):
    if text_file:
        word_file = text_to_word(text_file)
        st.download_button("Download Word Document", word_file)
